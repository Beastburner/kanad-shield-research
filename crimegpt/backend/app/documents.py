"""Stage 4 — GENERATION. Builds police legal documents as .docx.

Each document follows a consistent official layout:
  - a centered government / title / statute header block,
  - statutorily-grounded numbered fields (BNSS s.193 / s.187 / s.105-106 / s.190),
  - clean bordered tables for parties, property, witnesses,
  - signature / panch-witness blocks,
  - a disclaimer footer.

The *contents/fields* are grounded in the bare acts (see research/police-document-
formats.md). The exact Gujarat-Police proforma is State-prescribed — swap a real
.docx template per builder when available; the field structure stays the same."""

import uuid
from datetime import datetime, timezone
from pathlib import Path

from docx import Document as Docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from .config import settings
from .models import DISCLAIMER, DocType, ExtractedFacts, SuggestedSection

# (title, statutory citation) per document type
_DOC_META: dict[str, tuple[str, str]] = {
    "chargesheet": ("FINAL REPORT (CHARGE SHEET)",
                    "Under Section 193, Bharatiya Nagarik Suraksha Sanhita, 2023"),
    "remand_request": ("APPLICATION FOR POLICE CUSTODY REMAND",
                       "Under Section 187, Bharatiya Nagarik Suraksha Sanhita, 2023"),
    "seizure_receipt": ("SEIZURE MEMO / PANCHNAMA",
                        "Under Sections 103, 105 & 106, Bharatiya Nagarik Suraksha Sanhita, 2023"),
    "court_custody_letter": ("LETTER FOR JUDICIAL (COURT) CUSTODY",
                             "Under Section 190 r/w Section 187, Bharatiya Nagarik Suraksha Sanhita, 2023"),
    "accused_panchanama": ("ARREST PANCHNAMA",
                           "Under Section 105, Bharatiya Nagarik Suraksha Sanhita, 2023"),
    "medical_treatment_letter": ("REQUISITION FOR MEDICAL EXAMINATION / TREATMENT",
                                 "Under Section 51, Bharatiya Nagarik Suraksha Sanhita, 2023"),
    "face_identification_form": ("ACCUSED FACE IDENTIFICATION FORM",
                                 "Investigation aid — BNS / BNSS proceedings"),
    "lers_request": ("LAW ENFORCEMENT DATA REQUEST — META / WHATSAPP / INSTAGRAM (LERS)",
                     "Under Section 94 BNSS, 2023 r/w the IT Act, 2000 and platform LERS policy"),
}

_BLANK = "____________________"


# ---------------------------------------------------------------------------
# low-level formatting helpers
# ---------------------------------------------------------------------------
def _set_base_style(d: Docx) -> None:
    st = d.styles["Normal"]
    st.font.name = "Times New Roman"
    st.font.size = Pt(11)


def _center(d: Docx, text: str, size: int, bold=False, italic=False, color=None):
    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold, r.italic = bold, italic
    r.font.size = Pt(size)
    if color:
        r.font.color.rgb = color
    return p


def _rule(d: Docx):
    """A thin horizontal rule (paragraph bottom border)."""
    p = d.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "999999")
    pbdr.append(bottom)
    pPr.append(pbdr)
    return p


def _heading(d: Docx, text: str):
    p = d.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(11.5)
    return p


def _kv(d: Docx, pairs: list[tuple[str, str]]):
    t = d.add_table(rows=0, cols=2)
    t.style = "Table Grid"
    t.autofit = False
    for label, value in pairs:
        cells = t.add_row().cells
        cells[0].paragraphs[0].add_run(label).bold = True
        cells[1].text = str(value) if value not in (None, "") else _BLANK
        cells[0].width = Inches(2.3)
        cells[1].width = Inches(4.0)
    return t


def _table(d: Docx, headers: list[str], rows: list[list[str]]):
    t = d.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.paragraphs[0].add_run(h).bold = True
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = str(v)
    return t


def _para(d: Docx, text: str):
    return d.add_paragraph(text)


def _numbered(d: Docx, items: list[str]):
    for i, it in enumerate(items, 1):
        d.add_paragraph(f"{i}. {it}")


def _signoff(d: Docx, blocks: list[list[str]]):
    """Lay out signature blocks side by side (borderless table)."""
    d.add_paragraph()
    t = d.add_table(rows=1, cols=len(blocks))
    for i, lines in enumerate(blocks):
        cell = t.rows[0].cells[i]
        cell.paragraphs[0].add_run(_BLANK)
        for ln in lines:
            r = cell.add_paragraph().add_run(ln)
            r.font.size = Pt(9.5)


def _footer(d: Docx):
    _rule(d)
    _para_italic(d, DISCLAIMER, 9, RGBColor(0x99, 0x55, 0x00))
    ts = datetime.now(timezone.utc).isoformat(timespec="seconds")
    _para_italic(d, f"Generated (UTC): {ts}  ·  CrimeGPT v0.1 (AI-assisted draft)", 8,
                 RGBColor(0x80, 0x80, 0x80))


def _para_italic(d: Docx, text: str, size: int, color: RGBColor):
    p = d.add_paragraph()
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(size)
    r.font.color.rgb = color


# ---------------------------------------------------------------------------
# shared content blocks
# ---------------------------------------------------------------------------
def _title_block(d: Docx, doc_type: str):
    title, statute = _DOC_META[doc_type]
    _center(d, "GOVERNMENT OF GUJARAT", 10, bold=True)
    _center(d, "AHMEDABAD CITY POLICE  ·  CYBER CRIME BRANCH", 9)
    d.add_paragraph()
    _center(d, title, 15, bold=True)
    _center(d, statute, 9.5, italic=True)
    _rule(d)


def _case_header(d: Docx, case_number: str, facts: ExtractedFacts):
    fir_date = facts.dates[0] if facts.dates else ""
    today = datetime.now(timezone.utc).strftime("%d-%m-%Y")
    _kv(d, [
        ("District", "Ahmedabad"),
        ("Police Station", ""),
        ("FIR / Crime No.", case_number),
        ("FIR Date", fir_date),
        ("Document Date", today),
        ("Place of Occurrence", facts.location or ""),
    ])


def _sections_block(d: Docx, sections: list[SuggestedSection]):
    if not sections:
        _para(d, "No validated sections — officer review required.")
        return
    rows = [
        [s.code, s.section_no, s.heading or "", s.old_code_ref or "—"]
        for s in sections
    ]
    _table(d, ["Code", "Section", "Offence / Heading", "Old-code ref (cf.)"], rows)


def _property_block(d: Docx, facts: ExtractedFacts):
    if not facts.items:
        _para(d, "Nil.")
        return
    rows = [[str(i), it, "", "", ""] for i, it in enumerate(facts.items, 1)]
    _table(d, ["S.No", "Description", "Identifying marks", "Qty", "Est. value"], rows)


def _witness_block(d: Docx, facts: ExtractedFacts):
    rows = []
    if facts.complainant:
        rows.append(["1", facts.complainant, "Complainant"])
    rows.append([str(len(rows) + 1), _BLANK, "Eye-witness / Panch"])
    rows.append([str(len(rows) + 1), _BLANK, "Eye-witness / Panch"])
    _table(d, ["S.No", "Name & Address", "Category"], rows)


def _accused_block(d: Docx, facts: ExtractedFacts):
    accused = facts.accused or [""]
    for idx, name in enumerate(accused, 1):
        _para(d, f"Accused No. {idx}:").runs[0].bold = True
        _kv(d, [
            ("Name", name),
            ("Parentage (S/o, D/o, W/o)", ""),
            ("Address", ""),
            ("Age / Occupation", ""),
            ("Date & time of arrest", ""),
            ("Custody / Bail status", ""),
        ])


# ---------------------------------------------------------------------------
# per-document builders
# ---------------------------------------------------------------------------
def _build_chargesheet(d, facts, sections, case_number):
    _heading(d, "1. Parties")
    _kv(d, [("Complainant / Informant", facts.complainant or ""),
            ("Accused", ", ".join(facts.accused) or "")])

    _heading(d, "2. Nature of Information")
    _para(d, facts.events[0] if facts.events
          else "Cognizable offence reported vide the FIR referenced above.")

    _heading(d, "3. Sections Charged")
    _sections_block(d, sections)

    _heading(d, "4. Brief Facts of the Case (Purvani)")
    if facts.events:
        _numbered(d, facts.events)
    else:
        _para(d, _BLANK)

    _heading(d, "5. Property / Articles Seized")
    _property_block(d, facts)

    _heading(d, "6. Witnesses")
    _witness_block(d, facts)

    _heading(d, "7. Accused Details")
    _accused_block(d, facts)

    _heading(d, "8. Investigation Status")
    _kv(d, [
        ("Whether accused arrested", ""),
        ("Whether released on bond / bail bond", ""),
        ("Whether forwarded in custody (s.190)", ""),
        ("Chain of custody of electronic device (s.193(3))", ""),
        ("Whether investigation complete", ""),
    ])

    _signoff(d, [
        ["Investigating Officer", "Name & Rank, Police Station"],
        ["Station House Officer (endorsement)", "Name & Rank"],
    ])


def _build_remand_request(d, facts, sections, case_number):
    _center(d, "In the Court of the Judicial Magistrate First Class, Ahmedabad", 11, bold=True)
    d.add_paragraph()

    _heading(d, "Sections Invoked")
    _sections_block(d, sections)

    _heading(d, "Accused Produced")
    _accused_block(d, facts)

    _heading(d, "Custody Prayed For")
    _kv(d, [
        ("Type of custody", "Police Custody Remand (PCR)"),
        ("Number of days requested", ""),
        ("Date & time of arrest", ""),
        ("Date & time of production", ""),
        ("Earlier remand granted / days in custody", ""),
    ])

    _heading(d, "Grounds for Remand (custodial interrogation required for)")
    grounds = list(facts.events) + [
        "Recovery of case property and discovery of facts.",
        "Identification and confrontation of co-accused.",
    ]
    _numbered(d, grounds)

    _heading(d, "Enclosures & Case Diary")
    _para(d, "A copy of the FIR and the case diary maintained under Section 192 BNSS are "
             "produced/enclosed for the Hon'ble Court's perusal, demonstrating how the "
             "custody sought will assist the investigation.")

    _heading(d, "Prayer")
    _para(d, "It is therefore prayed that the Hon'ble Court be pleased to grant police "
             "custody of the accused for the period sought, in the interest of investigation.")

    _signoff(d, [["Investigating Officer", "Name & Rank, Police Station", "Date:"]])


def _build_seizure_receipt(d, facts, sections, case_number):
    _heading(d, "Sections Invoked")
    _sections_block(d, sections)

    _heading(d, "Seizure Details")
    _kv(d, [
        ("Date & time of seizure", ""),
        ("Place of seizure", facts.location or ""),
        ("Person from whom seized", facts.accused[0] if facts.accused else ""),
        ("Parentage / Address", ""),
    ])

    _heading(d, "Property Seized")
    if facts.items:
        rows = [[str(i), it, "", "", "", "", ""] for i, it in enumerate(facts.items, 1)]
        _table(d, ["S.No", "Description", "Make / Model", "IMEI / Serial No.",
                   "Distinguishing marks", "Qty", "Est. value"], rows)
    else:
        _para(d, "Nil.")

    _heading(d, "Electronic Evidence Integrity (BNSS s.193(3) / BSA s.63)")
    _kv(d, [
        ("SHA-256 hash of imaged data", ""),
        ("Hash of audio-video recording", ""),
        ("Recording forwarded to Magistrate (Y/N)", ""),
    ])
    _para(d, "Note: for any electronic device, the SHA-256 hash value MUST be recorded "
             "in this memo — absence of the hash may render the electronic evidence "
             "inadmissible.")

    _heading(d, "Packing & Sealing")
    _kv(d, [("Manner of packing / sealing", ""), ("Seal description", "")])

    _heading(d, "Recording (Section 105 BNSS)")
    _para(d, "The process of search and seizure, including preparation of this memo and "
             "obtaining witness signatures, was recorded by audio-video electronic means "
             "and the recording is being forwarded forthwith to the jurisdictional Magistrate.")

    _heading(d, "Panch Witnesses (two independent witnesses)")
    _table(d, ["S.No", "Name, Parentage & Address", "Signature"],
           [["1", _BLANK, ""], ["2", _BLANK, ""]])

    _signoff(d, [
        ["Signature of person from whom seized", "(acknowledgement)"],
        ["Seizing Officer", "Name & Rank, Police Station"],
    ])


def _build_court_custody_letter(d, facts, sections, case_number):
    _center(d, "In the Court of the Judicial Magistrate First Class, Ahmedabad", 11, bold=True)
    d.add_paragraph()

    _heading(d, "Sections Invoked")
    _sections_block(d, sections)

    _heading(d, "Accused Forwarded")
    _accused_block(d, facts)

    _heading(d, "Request")
    _para(d, "The accused named above is forwarded to the Hon'ble Court with a request to "
             "authorise judicial custody pending further investigation, on the following grounds:")
    _numbered(d, facts.events or [_BLANK])

    _signoff(d, [["Investigating Officer", "Name & Rank, Police Station", "Date:"]])


def _build_accused_panchanama(d, facts, sections, case_number):
    _heading(d, "Sections Invoked")
    _sections_block(d, sections)

    _heading(d, "Arrest Details")
    _kv(d, [
        ("Accused", ", ".join(facts.accused) or ""),
        ("Date & time of arrest", ""),
        ("Place of arrest", facts.location or ""),
    ])

    _heading(d, "Personal Search — Articles Found on Person")
    _property_block(d, facts)

    _heading(d, "Recording (Section 105 BNSS)")
    _para(d, "The arrest and personal search were conducted in the presence of two "
             "independent panch witnesses and recorded by audio-video electronic means.")

    _heading(d, "Panch Witnesses")
    _table(d, ["S.No", "Name, Parentage & Address", "Signature"],
           [["1", _BLANK, ""], ["2", _BLANK, ""]])

    _signoff(d, [["Arresting Officer", "Name & Rank, Police Station"]])


def _build_medical_treatment_letter(d, facts, sections, case_number):
    _para(d, "To,").runs[0].bold = True
    _para(d, "The Medical Officer / Superintendent,")
    _para(d, "________________________ Hospital, Ahmedabad.")

    _heading(d, "Subject")
    _para(d, "Requisition for medical examination and treatment under Section 51 BNSS, "
             "and issuance of a medico-legal certificate (MLC) for evidentiary purposes.")

    _heading(d, "Person to be Examined")
    person = facts.victims[0] if facts.victims else (facts.accused[0] if facts.accused else "")
    _kv(d, [
        ("Name", person),
        ("Whether victim / accused", ""),
        ("Age / Sex", ""),
        ("Nature of injuries / purpose", ""),
        ("Date & time of incident", ", ".join(facts.dates) or ""),
    ])

    _heading(d, "Sections Invoked")
    _sections_block(d, sections)

    _para(d, "You are requested to examine and treat the above person and furnish the MLC "
             "at the earliest in the interest of investigation.")

    _signoff(d, [["Investigating Officer", "Name & Rank, Police Station", "Date:"]])


def _build_face_identification_form(d, facts, sections, case_number):
    _heading(d, "Accused Particulars")
    _kv(d, [
        ("Name / Alias", ", ".join(facts.accused) or ""),
        ("Approximate age", ""),
        ("Height / Build", ""),
        ("Complexion", ""),
        ("Identifying marks (scar / tattoo / mole)", ""),
        ("Last seen location", facts.location or ""),
    ])

    _heading(d, "Photograph")
    _para(d, "[ Affix / attach accused photograph or captured image here ]")

    _heading(d, "Sections Invoked")
    _sections_block(d, sections)

    _heading(d, "Identified By (witness)")
    _table(d, ["Name & Address", "Relation to case", "Signature"],
           [[_BLANK, "", ""]])

    _signoff(d, [["Investigating Officer", "Name & Rank, Police Station"]])


def _build_lers_request(d, facts, sections, case_number):
    _para(d, "To,").runs[0].bold = True
    _para(d, "The Law Enforcement Response Team,")
    _para(d, "Meta Platforms, Inc. (Facebook / Instagram / WhatsApp)")
    _para(d, "via the Law Enforcement Online Request System (LERS).")

    _heading(d, "Subject")
    _para(d, "Request for preservation and disclosure of subscriber / account data "
             "in a criminal investigation, under Section 94 BNSS, 2023 read with the "
             "Information Technology Act, 2000 and the platform's law-enforcement policy.")

    _heading(d, "1. Requesting Authority")
    _kv(d, [
        ("Officer name & rank", ""),
        ("Police Station / Unit", ""),
        ("Official email (govt. domain)", ""),
        ("Contact number", ""),
    ])

    _heading(d, "2. Case Reference")
    _kv(d, [
        ("FIR / Crime No.", case_number),
        ("Date of incident", ", ".join(facts.dates) or ""),
        ("Nature of offence", facts.events[0] if facts.events else ""),
    ])

    _heading(d, "3. Sections Invoked")
    _sections_block(d, sections)

    _heading(d, "4. Target Account / Identifier(s)")
    _kv(d, [
        ("Platform (FB / IG / WhatsApp)", ""),
        ("Username / Profile URL", ""),
        ("Phone number (with country code)", ""),
        ("Email / Account ID", ""),
        ("Relevant date/time range (UTC)", ""),
    ])

    _heading(d, "5. Data Requested")
    _numbered(d, [
        "Basic subscriber information (name, registration details, phone/email).",
        "Account registration & login IP logs with timestamps (UTC).",
        "Non-content connection / message metadata for the period specified.",
        "Content data (only where accompanied by appropriate judicial process).",
    ])

    _heading(d, "6. Preservation Request (Emergency / Routine)")
    _para(d, "You are requested to PRESERVE all data associated with the above "
             "identifier(s) pending formal legal process, to prevent loss/deletion. "
             "Tick if this is an emergency disclosure request involving risk to life:  [  ]")

    _heading(d, "7. Legal Basis & Undertaking")
    _para(d, "This request is made for a bona fide criminal investigation. The data "
             "sought is necessary and proportionate to the offence under investigation. "
             "Content data, where requested, is supported by the requisite judicial order.")

    _signoff(d, [["Investigating Officer", "Name, Rank & Govt. Email", "Date:"],
                 ["Endorsing Officer (SHO)", "Name & Rank"]])


_BUILDERS = {
    "chargesheet": _build_chargesheet,
    "remand_request": _build_remand_request,
    "seizure_receipt": _build_seizure_receipt,
    "court_custody_letter": _build_court_custody_letter,
    "accused_panchanama": _build_accused_panchanama,
    "medical_treatment_letter": _build_medical_treatment_letter,
    "face_identification_form": _build_face_identification_form,
    "lers_request": _build_lers_request,
}


def generate(
    doc_type: DocType,
    case_number: str,
    facts: ExtractedFacts,
    sections: list[SuggestedSection],
) -> str:
    """Render a properly-formatted document to .docx and return its file path."""
    Path(settings.artifact_dir).mkdir(parents=True, exist_ok=True)
    d = Docx()
    _set_base_style(d)
    _title_block(d, doc_type)
    _case_header(d, case_number, facts)
    _BUILDERS[doc_type](d, facts, sections, case_number)
    _footer(d)

    fname = f"{doc_type}_{uuid.uuid4().hex[:8]}.docx"
    out_path = str(Path(settings.artifact_dir) / fname)
    d.save(out_path)
    return out_path

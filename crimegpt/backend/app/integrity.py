"""Evidence-integrity layer: SHA-256 hashing + BSA Section 63 (Part-A)
certificate auto-drafting. Every generated document passes through here."""

import hashlib
from datetime import datetime, timezone
from pathlib import Path

from docx import Document as Docx


def sha256_file(path: str) -> str:
    h = hashlib.sha256()
    with open(path, "rb") as f:
        for block in iter(lambda: f.read(8192), b""):
            h.update(block)
    return h.hexdigest()


def draft_s63_certificate(
    doc_path: str, sha256: str, case_number: str, doc_type: str, out_path: str
) -> str:
    """Auto-draft a BSA Section 63 Part-A certificate for an electronic record.

    BSA = Bharatiya Sakshya Adhiniyam, 2023. Section 63 governs admissibility of
    electronic records and requires a certificate. This is a DRAFT for officer
    review — it is not a self-certifying legal instrument."""
    ts = datetime.now(timezone.utc).isoformat()
    d = Docx()
    d.add_heading("CERTIFICATE UNDER SECTION 63", level=1)
    d.add_paragraph("Bharatiya Sakshya Adhiniyam, 2023 — Part-A (Electronic Record)")
    d.add_paragraph("")
    d.add_paragraph(
        "This certificate accompanies an electronic record produced by the "
        "CrimeGPT case-documentation system and is submitted for evidentiary "
        "purposes under Section 63 of the Bharatiya Sakshya Adhiniyam, 2023."
    )
    rows = [
        ("Case Number", case_number),
        ("Document Type", doc_type),
        ("Source File", Path(doc_path).name),
        ("SHA-256 Integrity Hash", sha256),
        ("Generated (UTC)", ts),
        ("Producing System", "CrimeGPT v0.1 (FastAPI / python-docx)"),
    ]
    table = d.add_table(rows=0, cols=2)
    table.style = "Light Grid Accent 1"
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = str(value)
    d.add_paragraph("")
    d.add_paragraph("Declaration under the four conditions of Section 63:")
    for cond in (
        "1. The computer/device producing this record was used regularly to "
        "store or process information during the relevant period.",
        "2. Information of the kind contained in this record was regularly fed "
        "into the device in the ordinary course of activities.",
        "3. The device was operating properly during the relevant period (or any "
        "malfunction did not affect the accuracy of the record).",
        "4. This record reproduces or is derived from information so stored, and "
        "its integrity is confirmed by the SHA-256 hash above.",
    ):
        d.add_paragraph(cond)
    d.add_paragraph("")
    d.add_paragraph(
        "Particulars to be completed by the certifying officer (Part-A):"
    )
    for line in (
        "Name of certifying person: ____________________",
        "Designation: ____________________",
        "Date & place: ____________________",
        "Signature: ____________________",
    ):
        d.add_paragraph(line)
    d.add_paragraph("")
    d.add_heading("Part-B — to be completed by the Expert", level=2)
    d.add_paragraph(
        "Left blank by design. Part-B (technical/forensic certification, "
        "independent SHA-256 re-validation) must be completed by an Examiner of "
        "Electronic Evidence / cyber-forensic expert — the system does not "
        "auto-assert expert findings."
    )
    for line in (
        "Expert name & qualifications: ____________________",
        "Forensic tools / environment: ____________________",
        "Independently re-computed SHA-256: ____________________",
        "Findings, date, signature & lab seal: ____________________",
    ):
        d.add_paragraph(line)
    d.add_paragraph("")
    d.add_paragraph(
        "AI-assisted draft — officer review required. Verify all particulars "
        "before tendering as evidence."
    ).italic = True
    d.save(out_path)
    return out_path

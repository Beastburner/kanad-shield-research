"""Multilingual I/O layer (PS functional requirement: Gujarati / Hindi / English).

Translation uses the same Groq LLM, so there is no extra dependency. The FIR
narrative may also be submitted directly in Hindi/Gujarati — the extraction agent
handles multilingual input natively; this module also translates UI text and
generated .docx documents into the officer's language."""

import json

from pydantic import BaseModel

from .pipeline import llm

LANGS = {"en": "English", "hi": "Hindi", "gu": "Gujarati"}


async def translate(text: str, target: str, source: str = "auto") -> str:
    if target not in LANGS:
        raise ValueError(f"unsupported target language: {target}")
    if not text.strip():
        return text
    src = LANGS.get(source, "the source language")
    system = (
        f"You are a precise legal translator. Translate the user's text from {src} "
        f"into {LANGS[target]}. Preserve legal section numbers (e.g. BNS 303), "
        "names, dates and numbers exactly. Output ONLY the translation, no notes."
    )
    return await llm.complete_text(system, text)


# --- document translation -------------------------------------------------
class _Batch(BaseModel):
    # index (as string) -> translated text; index-keyed so a dropped item simply
    # falls back to the original English rather than breaking the whole batch.
    translations: dict[str, str]


async def _translate_list(texts: list[str], target: str) -> list[str]:
    payload = {str(i): t for i, t in enumerate(texts)}
    system = (
        f"You are a legal translator for Indian police documents. The input is a JSON "
        f"object mapping numeric keys to English strings. Return a JSON object with key "
        f"'translations' mapping the SAME keys to each string translated into "
        f"{LANGS[target]}. Keep EVERY key; translate every value. Translate only the "
        "natural-language words. Keep numbers, dates, FIR/case numbers, SHA-256 hashes, "
        "and runs of underscores unchanged; you may render law citations naturally "
        "(e.g. 'Section 193' as is idiomatic in the target language)."
    )
    out = await llm.complete_json(system, json.dumps(payload, ensure_ascii=False), _Batch)
    m = out.translations
    # map back by index; any missing/extra key falls back to the original text
    return [m.get(str(i), texts[i]) for i in range(len(texts))]


def _translatable(text: str) -> bool:
    t = text.strip()
    return bool(t) and any(c.isalpha() for c in t) and set(t) != {"_"}


def _set_paragraph_text(par, text: str) -> None:
    if par.runs:
        par.runs[0].text = text
        for r in par.runs[1:]:
            r.text = ""
    else:
        par.add_run(text)


async def translate_docx_file(path: str, target: str) -> None:
    """Translate a generated .docx in place into `target` (hi/gu). On any failure
    the document is left in English (translation is best-effort, never fatal)."""
    if target not in ("hi", "gu"):
        return
    from docx import Document

    doc = Document(path)
    pars = [p for p in _iter_paragraphs(doc) if _translatable(p.text)]
    if not pars:
        return
    try:
        translated = await _translate_list([p.text for p in pars], target)
    except Exception as e:  # LLM/parse error -> leave English
        print(f"[translate] doc translation failed ({e}); leaving English")
        return
    if len(translated) != len(pars):  # count mismatch -> leave English
        print("[translate] translation count mismatch; leaving English")
        return
    for par, new_text in zip(pars, translated):
        _set_paragraph_text(par, new_text)
    doc.save(path)


def _iter_paragraphs(doc):
    yield from doc.paragraphs
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                yield from cell.paragraphs
